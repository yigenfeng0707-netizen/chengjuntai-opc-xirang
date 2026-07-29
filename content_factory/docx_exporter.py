# -*- coding: utf-8 -*-
"""Markdown → 富文本 Word (.docx)：标题 / 粗体 / 列表"""
import os
import re
from typing import Optional

from config_loader import ARTICLES_DIR, DATA_DIR, ROOT
import op_logger

EXPORT_DOCX_DIR = os.path.join(ROOT, "export_docx")
ARTICLES_META = os.path.join(DATA_DIR, "articles_meta.json")
os.makedirs(EXPORT_DOCX_DIR, exist_ok=True)

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_CODE_RE = re.compile(r"`([^`]+)`")


def _add_runs(paragraph, text: str):
    """把含 **粗体** / *斜体* / `代码` 的行拆成 runs。"""
    if not text:
        return
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^*]+\*|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_text: str, out_path: str, title: Optional[str] = None) -> str:
    """简易 Markdown → .docx，返回绝对路径。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    if title:
        doc.add_heading(title, level=0)

    # 去掉 YAML front matter
    md_text = re.sub(r"^---\n.*?\n---\n", "", md_text or "", flags=re.S)
    in_code = False
    code_buf = []

    for raw in (md_text or "").split("\n"):
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^[-*]\s+", "", line))
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("|") and "|" in line[1:]:
            # 简易表格行：拼成正文，避免复杂表格布局
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^:?-+:?$", c or "") for c in cells):
                continue
            p = doc.add_paragraph()
            _add_runs(p, " | ".join(cells))
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    op_logger.log("docx_exporter", f"已导出 Word: {out_path}")
    return os.path.abspath(out_path)


def export_markdown_file(md_path: str, out_name: Optional[str] = None, title: Optional[str] = None) -> str:
    with open(md_path, "r", encoding="utf-8") as f:
        md = f.read()
    name = out_name or (os.path.splitext(os.path.basename(md_path))[0] + ".docx")
    if not name.lower().endswith(".docx"):
        name += ".docx"
    out = os.path.join(EXPORT_DOCX_DIR, name)
    return md_to_docx(md, out, title=title)


def export_article_by_id(article_id: str) -> str:
    """按 articles_meta / 文件名导出文章 Word。"""
    import json
    file_name = None
    title = article_id
    if os.path.exists(ARTICLES_META):
        with open(ARTICLES_META, "r", encoding="utf-8") as f:
            metas = json.load(f)
        for m in metas:
            if m.get("id") == article_id:
                file_name = m.get("file")
                title = m.get("title") or title
                break
    if not file_name:
        for fn in os.listdir(ARTICLES_DIR):
            if fn.startswith(article_id) and fn.endswith(".md"):
                file_name = fn
                break
    if not file_name:
        raise FileNotFoundError(f"文章不存在: {article_id}")
    md_path = file_name if os.path.isabs(file_name) else os.path.join(ARTICLES_DIR, file_name)
    if not os.path.exists(md_path):
        raise FileNotFoundError(md_path)
    return export_markdown_file(md_path, out_name=f"{article_id}.docx", title=title)
